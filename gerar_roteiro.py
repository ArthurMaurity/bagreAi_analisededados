"""
Gera roteiro_bagre.docx — pronto para importar no Google Docs.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── Estilos base ──────────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

def titulo(texto, nivel=1):
    p = doc.add_heading(texto, level=nivel)
    return p

def paragrafo(texto, negrito=False, italico=False, cor=None, tamanho=11, alinhamento=None):
    p = doc.add_paragraph()
    if alinhamento:
        p.alignment = alinhamento
    run = p.add_run(texto)
    run.bold = negrito
    run.italic = italico
    run.font.size = Pt(tamanho)
    if cor:
        r, g, b = cor
        run.font.color.rgb = RGBColor(r, g, b)
    return p

def linha():
    doc.add_paragraph('─' * 70)

def bloco(tempo, titulo_bloco, telas, narracoes):
    """Renderiza um bloco do roteiro."""
    p = doc.add_paragraph()
    run = p.add_run(f"[{tempo}]  {titulo_bloco}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x7E, 0x34)

    for item in telas:
        p2 = doc.add_paragraph(style='List Bullet')
        run2 = p2.add_run(item)
        run2.italic = True
        run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        run2.font.size = Pt(10)

    for fala in narracoes:
        p3 = doc.add_paragraph()
        p3.paragraph_format.left_indent = Inches(0.3)
        run3 = p3.add_run(f'"{fala}"')
        run3.font.size = Pt(11)

    doc.add_paragraph()

# ── CAPA ──────────────────────────────────────────────────────────────────────
p_titulo = doc.add_paragraph()
p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_titulo.add_run('BAGRE.AI')
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1F, 0x7E, 0x34)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run('Roteiro — Vídeo Tutorial (5 minutos)')
r2.font.size = Pt(16)
r2.bold = True

p_desc = doc.add_paragraph()
p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p_desc.add_run('Demonstração completa de todas as features para apresentação ao professor')
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()

# ── ESTRUTURA DE TEMPO ────────────────────────────────────────────────────────
titulo('Estrutura de Tempo', nivel=1)

tabela = doc.add_table(rows=1, cols=3)
tabela.style = 'Table Grid'
cab = tabela.rows[0].cells
for i, txt in enumerate(['Bloco', 'Cena', 'Tempo']):
    cab[i].paragraphs[0].add_run(txt).bold = True

blocos_tabela = [
    ('1', 'Abertura + conceito',           '0:00 – 0:30'),
    ('2', 'Tela de planos + login',         '0:30 – 0:50'),
    ('3', 'Dashboard + Pedigree da Semana', '0:50 – 1:15'),
    ('4', 'Scout de Jogador (Várzea)',      '1:15 – 1:55'),
    ('5', 'Hype Index (Olheiro)',           '1:55 – 2:35'),
    ('6', 'Espelho Pedigree (Olheiro)',     '2:35 – 3:10'),
    ('7', 'Ciclo de Vida (Diretor)',        '3:10 – 3:40'),
    ('8', 'Golden List (Diretor)',          '3:40 – 4:05'),
    ('9', 'Relatório DOCX (Diretor)',       '4:05 – 4:35'),
    ('10','Encerramento',                   '4:35 – 5:00'),
]
for num, cena, tempo in blocos_tabela:
    row = tabela.add_row().cells
    row[0].paragraphs[0].add_run(num)
    row[1].paragraphs[0].add_run(cena)
    row[2].paragraphs[0].add_run(tempo)

doc.add_paragraph()
doc.add_page_break()

# ── ROTEIRO ───────────────────────────────────────────────────────────────────
titulo('Roteiro Completo', nivel=1)
doc.add_paragraph()

bloco(
    '0:00 – 0:30', 'BLOCO 1 — ABERTURA',
    ['TELA: Logo Bagre.ai em fundo preto, trilha suave ao fundo'],
    [
        'O mercado de transferências do futebol movimenta bilhões de euros por ano — '
        'e a maior parte das decisões ainda é tomada no achismo.',
        'O Bagre.ai muda isso. Somos uma plataforma de inteligência de mercado que usa '
        'dados reais para identificar dois perfis: o Pedigree — jogador subvalorizado, '
        'alto rendimento por baixo custo — e o PediRato — caro demais para o que entrega.',
        'Vamos ver como funciona.',
    ]
)

bloco(
    '0:30 – 0:50', 'BLOCO 2 — TELA DE PLANOS',
    [
        'TELA: Browser abre em localhost:8000 — tela de seleção de planos',
        'AÇÃO: Clica no card "Diretor"',
    ],
    [
        'Ao acessar o sistema, o usuário escolhe seu plano diretamente na tela inicial. '
        'São três níveis: Várzea, gratuito — acesso básico. '
        'Olheiro, para analistas — R$ 29,90 por mês. '
        'Diretor, para clubes e agentes — R$ 997 por mês, acesso total.',
        'Vou entrar como Diretor para mostrar todas as funcionalidades.',
    ]
)

bloco(
    '0:50 – 1:15', 'BLOCO 3 — DASHBOARD',
    ['TELA: Dashboard carrega com card "Pedigree da Semana" em destaque'],
    [
        'O Dashboard é o painel central. No topo, temos o Pedigree da Semana — '
        'um destaque automático que muda todo domingo, sempre apontando o jogador '
        'com melhor relação custo-benefício do nosso banco de dados.',
        'Aqui vemos nome, clube, gols, assistências, valor de mercado e o Índice PediRato — '
        'nossa métrica principal: Gols mais Assistências dividido pelo Valor de Mercado em milhões. '
        'Quanto maior, mais eficiente financeiramente.',
    ]
)

bloco(
    '1:15 – 1:55', 'BLOCO 4 — SCOUT DE JOGADOR',
    [
        'AÇÃO: Clica em "Scout de Jogador" no menu lateral',
        'AÇÃO: Digita "Arrascaeta" no campo de busca e pressiona Buscar',
        'TELA: Resultado com Arrascaeta — Pedigree, pedirato 3.14',
    ],
    [
        'A função central do sistema é o Scout de Jogador. '
        'Digito o nome de um atleta e o sistema dispara um waterfall de fontes em paralelo: '
        'primeiro o cache interno, depois a API-Football, Transfermarkt e FBref.',
        'Em segundos temos gols, assistências, valor de mercado, clube, liga e a classificação automática.',
        'Arrascaeta — 25 gols, 19 assistências, 14 milhões de euros. '
        'Índice PediRato de 3.14 — classificado como Pedigree. '
        'Em outras palavras: entrega mais do que o preço sugere.',
    ]
)

bloco(
    '1:55 – 2:35', 'BLOCO 5 — HYPE INDEX',
    [
        'AÇÃO: Clica em "Hype Index" no menu',
        'TELA: Campos preenchidos com Arrascaeta, Vinicius Jr., Mbappé, Griezmann, Haaland',
        'AÇÃO: Clica em "Analisar"',
        'TELA: Gráfico scatter com reta de regressão — vermelho PediRato, verde Pedigree',
    ],
    [
        'O Hype Index usa regressão linear para cruzar performance real com valor de mercado. '
        'Informo uma lista de jogadores e o modelo calcula o valor esperado de cada um '
        'com base na reta de regressão.',
        'Jogadores com resíduo acima de um desvio padrão são PediRatos — pagamos caro demais. '
        'Abaixo de um desvio, são Pedigrees — oportunidade de mercado.',
        'No gráfico, vermelho é PediRato, verde é Pedigree. '
        'Veja Mbappé — acima da reta, caro para o que entrega na temporada atual. '
        'Griezmann — abaixo, entregando mais do que o preço sugere.',
    ]
)

bloco(
    '2:35 – 3:10', 'BLOCO 6 — ESPELHO PEDIGREE',
    [
        'AÇÃO: Clica em "Espelho Pedigree" no menu',
        'TELA: Referencial = Mbappé (€250M), pool com vários jogadores',
        'AÇÃO: Clica em "Analisar"',
        'TELA: Gráfico radar — Mbappé em vermelho, espelhos em verde',
    ],
    [
        'O Espelho Pedigree responde uma pergunta prática dos diretores: '
        '"Quero contratar alguém parecido com esse craque — mas sem gastar o mesmo."',
        'O sistema calcula similaridade cosseno entre os vetores de performance, '
        'filtra quem custa menos de 60% do referencial e retorna os 3 mais similares '
        'com a economia estimada.',
        'Griezmann, 97% de similaridade com Mbappé — com uma economia de 230 milhões de euros. '
        'Esse é o poder do Bagre.ai.',
    ]
)

bloco(
    '3:10 – 3:40', 'BLOCO 7 — CICLO DE VIDA',
    [
        'AÇÃO: Clica em "Ciclo de Vida" no menu',
        'TELA: Lista com Endrick 18a, Raphinha 28a, Griezmann 33a, Haaland 24a, Modric 39a',
        'AÇÃO: Clica em "Analisar"',
        'TELA: Curva parabólica com pico destacado e janela de revenda sombreada em verde',
    ],
    [
        'O Ciclo de Vida usa regressão polinomial de grau 2 para modelar a curva de valor '
        'por idade de um conjunto de jogadores.',
        'O modelo identifica o pico de valor estimado e a janela ideal de revenda — '
        'o período em que o jogador ainda está acima de 80% do seu valor máximo.',
        'Pico estimado aos 25 anos e meio, janela de revenda entre 20 e 31 anos. '
        'Informação direta para decisões de contratação e timing de venda.',
    ]
)

bloco(
    '3:40 – 4:05', 'BLOCO 8 — GOLDEN LIST',
    [
        'AÇÃO: Clica em "Golden List" no menu',
        'AÇÃO: Clica em "Usar Dataset Padrão" e depois "Gerar Golden List"',
        'TELA: Tabela ranqueada com posição, nome, gols, assists, valor e pedirato',
    ],
    [
        'A Golden List é o ranking dos melhores Pedigrees do mercado. '
        'Posso usar o dataset curado do Bagre.ai ou informar meu próprio pool.',
        'O sistema filtra automaticamente quem está acima da média de eficiência do pool '
        'e ordena por índice PediRato decrescente.',
        'Essa é a lista de compras do diretor esportivo.',
    ]
)

bloco(
    '4:05 – 4:35', 'BLOCO 9 — RELATÓRIO DOCX',
    [
        'AÇÃO: Clica em "Relatório DOCX" no menu lateral',
        'AÇÃO: Clica em "Gerar Relatório .DOCX"',
        'AÇÃO: Clica no link de download — documento Word abre',
        'TELA: Capa Bagre.ai com tabela comparativa e gráficos',
    ],
    [
        'Por fim, o plano Diretor gera um relatório Word automatizado com todos os dados '
        'da sessão — tabela comparativa, gráficos de hype, radar de espelho pedigree, '
        'curva de ciclo de vida e metodologia.',
        'Relatório pronto para enviar ao conselho do clube ou ao agente do atleta.',
    ]
)

bloco(
    '4:35 – 5:00', 'BLOCO 10 — ENCERRAMENTO',
    [
        'TELA: Dashboard — Logo Bagre.ai em destaque',
        'TELA: Fade para preto com tagline "Inteligência de Mercado no Futebol"',
    ],
    [
        'Bagre.ai — dados reais, decisões melhores.',
        'Do scout gratuito ao relatório executivo, tudo numa plataforma '
        'construída do zero para o mercado de transferências do futebol.',
        'Obrigado.',
    ]
)

doc.add_page_break()

# ── DICAS DE GRAVAÇÃO ─────────────────────────────────────────────────────────
titulo('Dicas de Gravação', nivel=1)

dicas = [
    ('Resolução', '1920×1080 mínimo'),
    ('Browser', 'Zoom em 90% para caber mais conteúdo na tela'),
    ('Dados', 'Deixe os campos pré-preenchidos antes de gravar cada bloco'),
    ('Pausas', 'Após cada gráfico carregar, pause 2s antes de continuar a narração'),
    ('Edição', 'Corte os tempos de carregamento da API — scout pode demorar 3–5s'),
    ('Trilha', 'Suave, sem letra, volume baixo — não compete com a narração'),
]

for chave, valor in dicas:
    p = doc.add_paragraph(style='List Bullet')
    run_k = p.add_run(f'{chave}: ')
    run_k.bold = True
    p.add_run(valor)

# ── SALVAR ────────────────────────────────────────────────────────────────────
doc.save('roteiro_bagre.docx')
print('roteiro_bagre.docx gerado com sucesso.')
