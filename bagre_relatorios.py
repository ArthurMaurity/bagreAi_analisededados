import matplotlib
matplotlib.use('Agg') # Evita problemas com falta de display gráfico no terminal
import matplotlib.pyplot as plt
import seaborn as sns
import pandas as pd
import os
from datetime import datetime

class GeradorRelatorios:
    """
    Gerador de relatórios e visualizações para o Bagre.ai.
    Cria tabelas em Markdown e gráficos usando Matplotlib/Seaborn.
    """
    
    @staticmethod
    def gerar_relatorio_markdown(historico_jogadores, filename="relatorio_scout.md"):
        """
        Gera um arquivo Markdown formatado com a tabela dos jogadores pesquisados.
        Retorna o texto do Markdown gerado ou None se houver erro.
        """
        if not historico_jogadores:
            print("Nenhum dado no histórico para gerar o relatório.")
            return None
            
        try:
            # Garantir que o diretório pai existe
            dir_name = os.path.dirname(filename)
            if dir_name:
                os.makedirs(dir_name, exist_ok=True)
                
            md_lines = []
            md_lines.append("# Relatório de Scouting - Bagre.ai\n")
            md_lines.append("Este relatório contém os dados consolidados dos atletas analisados.\n")
            md_lines.append("## Tabela Comparativa\n")
            md_lines.append("| Atleta | Gols | Assistências | Valor de Mercado | Índice PediRato |")
            md_lines.append("| :--- | :---: | :---: | :---: | :---: |")
            
            for jogador in historico_jogadores:
                nome = jogador.get("nome", "Desconhecido").upper()
                gols = jogador.get("gols", 0)
                assists = jogador.get("assists", 0)
                valor = jogador.get("valor_milhoes", 0.0)
                try:
                    gols = int(gols) if gols not in (None, '—', '') else 0
                except: gols = 0
                try:
                    assists = int(assists) if assists not in (None, '—', '') else 0
                except: assists = 0
                try:
                    valor = float(valor) if valor not in (None, '—', '') else 0.0
                except: valor = 0.0

                pedirato = jogador.get("pedirato") or jogador.get("pedirato_score") or jogador.get("score") or 0.0
                try:
                    pedirato = float(pedirato) if pedirato not in (None, '—', '') else 0.0
                except: pedirato = 0.0
                if pedirato == 0.0 and valor > 0.0:
                    pedirato = (gols + assists) / valor
                
                md_lines.append(f"| **{nome}** | {gols} | {assists} | €{valor}M | **{pedirato:.2f}** |")
                
            md_lines.append("\n*O Índice PediRato é calculado como: (Gols + Assistências) / Valor de Mercado (em milhões de euros). Quanto maior o índice, melhor a eficiência financeira/esportiva.*\n")
            
            markdown_content = "\n".join(md_lines)
            
            # Salvar no arquivo
            with open(filename, "w", encoding="utf-8") as f:
                f.write(markdown_content)
                
            print(f"Relatório Markdown salvo com sucesso em: {filename}")
            return markdown_content
        except Exception as e:
            print(f"Erro ao gerar relatório Markdown: {e}")
            return None

    @staticmethod
    def gerar_grafico_comparativo(historico_jogadores, filename="comparativo_pedirato.png"):
        """
        Gera um gráfico de barras comparando o Índice PediRato dos jogadores da sessão.
        Retorna o caminho do arquivo gerado ou None se houver erro.
        """
        if not historico_jogadores:
            print("Nenhum dado no histórico para gerar o gráfico.")
            return None
            
        try:
            # Garantir que o diretório pai existe
            dir_name = os.path.dirname(filename)
            if dir_name:
                os.makedirs(dir_name, exist_ok=True)
                
            # Converter para DataFrame
            dados = []
            for j in historico_jogadores:
                gols = j.get("gols", 0)
                assists = j.get("assists", 0)
                valor = j.get("valor_milhoes", 0.0)
                try:
                    gols = int(gols) if gols not in (None, '—', '') else 0
                except: gols = 0
                try:
                    assists = int(assists) if assists not in (None, '—', '') else 0
                except: assists = 0
                try:
                    valor = float(valor) if valor not in (None, '—', '') else 0.0
                except: valor = 0.0

                pedirato = j.get("pedirato") or j.get("pedirato_score") or j.get("score") or 0.0
                try:
                    pedirato = float(pedirato) if pedirato not in (None, '—', '') else 0.0
                except: pedirato = 0.0
                if pedirato == 0.0 and valor > 0.0:
                    pedirato = (gols + assists) / valor

                dados.append({
                    "Nome": j.get("nome", "Desconhecido").upper(),
                    "PediRato": pedirato
                })
            
            df = pd.DataFrame(dados)
            # Ordenar por PediRato decrescente
            df = df.sort_values(by="PediRato", ascending=False)
            
            # Configurar o estilo do Seaborn para uma estética premium
            sns.set_theme(style="darkgrid")
            plt.figure(figsize=(10, 6))
            
            # Paleta de cores moderna (degradê)
            colors = sns.color_palette("viridis", len(df))
            
            # Criar gráfico de barras
            ax = sns.barplot(
                x="PediRato", 
                y="Nome", 
                data=df, 
                palette=colors,
                hue="Nome",
                legend=False
            )
            
            # Adicionar rótulos de dados nas barras
            for i, p in enumerate(ax.patches):
                width = p.get_width()
                ax.text(
                    width + 0.05, 
                    p.get_y() + p.get_height() / 2, 
                    f"{width:.2f}", 
                    ha="left", 
                    va="center",
                    fontweight="bold"
                )
                
            plt.title("Comparação de Eficiência - Índice PediRato", fontsize=16, fontweight="bold", pad=15)
            plt.xlabel("Índice PediRato (Eficiência)", fontsize=12, fontweight="bold")
            plt.ylabel("Atleta", fontsize=12, fontweight="bold")
            
            plt.tight_layout()
            plt.savefig(filename, dpi=300)
            plt.close()
            
            print(f"Gráfico de comparação salvo com sucesso em: {filename}")
            return filename
        except Exception as e:
            print(f"Erro ao gerar gráfico comparativo: {e}")
            return None

    @staticmethod
    def gerar_relatorio_completo_docx(historico_jogadores, resultados_analytics=None, filename="relatorio_bagre_completo.docx"):
        """
        Gera relatório Word completo (.docx) com capa, tabela, gráficos e metodologia.
        Requer python-docx. Retorna o caminho do arquivo ou None em caso de erro.
        """
        if not historico_jogadores:
            print("Nenhum dado no histórico para gerar o relatório.")
            return None

        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_ALIGN_VERTICAL
            from docx.oxml.shared import OxmlElement
            from docx.oxml.ns import qn

            STATIC_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static")

            def _sombrear_celula(celula, hex_fill):
                tc = celula._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"),   "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"),  hex_fill)
                tcPr.append(shd)

            def _run_colorido(paragrafo, texto, hex_rgb, tamanho=None, negrito=False):
                r = hex_rgb.lstrip("#")
                run = paragrafo.add_run(texto)
                run.font.color.rgb = RGBColor(int(r[0:2], 16), int(r[2:4], 16), int(r[4:6], 16))
                run.bold = negrito
                if tamanho:
                    run.font.size = Pt(tamanho)
                return run

            def _classificar(pedirato):
                if pedirato > 1.5:
                    return "PEDIGREE"
                if pedirato < 0.5:
                    return "PEDIRATO"
                return "REGULAR"

            doc = Document()

            # ── Rodapé em todas as páginas ─────────────────────────────────────
            secao = doc.sections[0]
            rodape = secao.footer.paragraphs[0]
            rodape.text = "BAGRE.AI | Confidencial - Tier Diretor | Documento gerado automaticamente"
            rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rodape.runs[0].font.size = Pt(8)
            rodape.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            # ── 1. CAPA ────────────────────────────────────────────────────────
            p_titulo = doc.add_paragraph()
            p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run_colorido(p_titulo, "BAGRE.AI", "39FF14", tamanho=32, negrito=True)

            p_sub = doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run_colorido(p_sub, "RELATORIO DE INTELIGENCIA DE MERCADO", "1A1A1A", tamanho=18, negrito=True)

            p_sub2 = doc.add_paragraph()
            p_sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run_colorido(p_sub2, "Analise de Eficiencia Financeira e Scouting Avancado", "444444", tamanho=13)

            doc.add_paragraph()

            p_data = doc.add_paragraph()
            p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
            data_str = datetime.now().strftime("%d/%m/%Y %H:%M")
            _run_colorido(p_data, f"Gerado em: {data_str}", "888888", tamanho=10)

            doc.add_paragraph("_" * 80)
            doc.add_page_break()

            # ── 2. SUMÁRIO EXECUTIVO ───────────────────────────────────────────
            doc.add_heading("1. Sumario Executivo", level=1)
            doc.add_paragraph(
                "O Indice PediRato e uma metrica proprietaria do Bagre.ai que mede a eficiencia "
                "financeira de um atleta de futebol. Calculado como a razao entre contribuicoes "
                "ofensivas (gols + assistencias) e o valor de mercado em milhoes de euros, "
                "o indice identifica jogadores subvalorizados (Pedigrees) e superestimados (PediRatos).\n\n"
                "Formula: PediRato = (Gols + Assistencias) / Valor de Mercado (EM)\n\n"
                "Quanto maior o indice, melhor a eficiencia financeira do atleta."
            )
            doc.add_paragraph(f"Total de atletas analisados nesta sessao: {len(historico_jogadores)}")

            # ── 3. TABELA COMPARATIVA ──────────────────────────────────────────
            doc.add_heading("2. Tabela Comparativa de Atletas", level=1)

            colunas = ["Atleta", "Gols", "Assists", "Valor (EM)", "PediRato", "Classificacao"]
            tabela = doc.add_table(rows=1, cols=len(colunas))
            tabela.style = "Table Grid"

            # Cabeçalho com fundo cinza escuro
            cab = tabela.rows[0].cells
            for i, col in enumerate(colunas):
                _sombrear_celula(cab[i], "444444")
                p = cab[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(col)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)

            for j in historico_jogadores:
                linha = tabela.add_row().cells
                gols = j.get("gols", 0)
                assists = j.get("assists", 0)
                valor = j.get("valor_milhoes", 0.0)
                try:
                    gols = int(gols) if gols not in (None, '—', '') else 0
                except: gols = 0
                try:
                    assists = int(assists) if assists not in (None, '—', '') else 0
                except: assists = 0
                try:
                    valor = float(valor) if valor not in (None, '—', '') else 0.0
                except: valor = 0.0

                pedirato = j.get("pedirato") or j.get("pedirato_score") or j.get("score") or 0.0
                try:
                    pedirato = float(pedirato) if pedirato not in (None, '—', '') else 0.0
                except: pedirato = 0.0
                if pedirato == 0.0 and valor > 0.0:
                    pedirato = (gols + assists) / valor

                cls   = _classificar(pedirato)
                dados = [
                    j.get("nome", "?").upper(),
                    str(gols),
                    str(assists),
                    f"{valor:.2f}",
                    f"{pedirato:.3f}",
                    cls,
                ]
                cor_cls = {"PEDIGREE": "1A6B1A", "PEDIRATO": "8B0000", "REGULAR": "333333"}
                for i, texto in enumerate(dados):
                    p = linha[i].paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(texto)
                    run.font.size = Pt(10)
                    if i == 5:  # coluna Classificação
                        run.bold = True
                        run.font.color.rgb = RGBColor(
                            *[int(cor_cls[cls][k:k+2], 16) for k in (0, 2, 4)]
                        )

            # ── GERAÇÃO DINÂMICA DOS GRÁFICOS ────────────────────────────────
            caminho_comparativo = filename.replace(".docx", "_comparativo.png")
            caminho_hype = filename.replace(".docx", "_hype.png")
            caminho_radar = filename.replace(".docx", "_radar.png")
            caminho_ciclo = filename.replace(".docx", "_ciclo.png")

            # 1. Comparativo
            GeradorRelatorios.gerar_grafico_comparativo(historico_jogadores, caminho_comparativo)

            # 2. Hype Index (requer mín. 3 jogadores)
            if len(historico_jogadores) >= 3:
                try:
                    from bagre_analytics import analisar_hype_index
                    df_hype = pd.DataFrame(historico_jogadores)
                    # Certificar que colunas necessárias existem
                    if "valor_milhoes" in df_hype.columns and "gols" in df_hype.columns and "assists" in df_hype.columns:
                        analisar_hype_index(df_hype, img_path=caminho_hype)
                except Exception as ex:
                    print(f"Erro ao gerar hype_index dinamico: {ex}")

            # 3. Radar (requer mín. 2 jogadores)
            if len(historico_jogadores) >= 2:
                try:
                    from bagre_analytics import espelho_pedigree
                    pool_restante = sorted(historico_jogadores, key=lambda x: x.get("valor_milhoes", 0), reverse=True)
                    jogador_ref = pool_restante[0]
                    pool = pool_restante[1:]
                    espelho_pedigree(jogador_ref, pool, img_path=caminho_radar)
                except Exception as ex:
                    print(f"Erro ao gerar radar dinamico: {ex}")

            # 4. Ciclo de vida (requer mín. 3 jogadores com idades válidas e distintas)
            try:
                from bagre_analytics import analisar_ciclo_de_vida
                jogadores_com_idade = []
                for j in historico_jogadores:
                    idade = j.get("idade")
                    # Se idade for string de número ou int, inclui
                    if idade is not None and str(idade).isdigit():
                        j_copy = dict(j)
                        j_copy["idade"] = int(idade)
                        jogadores_com_idade.append(j_copy)
                
                if len(jogadores_com_idade) >= 3:
                    df_ciclo = pd.DataFrame(jogadores_com_idade)
                    if df_ciclo["idade"].nunique() >= 2:
                        analisar_ciclo_de_vida(df_ciclo, img_path=caminho_ciclo)
            except Exception as ex:
                print(f"Erro ao gerar ciclo_vida dinamico: {ex}")

            # ── 4. GRÁFICOS ANALÍTICOS ─────────────────────────────────────────
            doc.add_page_break()
            doc.add_heading("3. Graficos Analiticos", level=1)

            graficos_dinamicos = [
                (caminho_comparativo, "Grafico 1: Comparativo de Eficiencia (Indice PediRato)"),
                (caminho_hype,        "Grafico 2: Hype Index - Regressao Linear"),
                (caminho_radar,       "Grafico 3: Espelho Pedigree - Grafico Radar"),
                (caminho_ciclo,       "Grafico 4: Curva de Ciclo de Vida por Idade"),
            ]

            inseriu = False
            for caminho_img, legenda in graficos_dinamicos:
                if os.path.exists(caminho_img):
                    p_leg = doc.add_paragraph()
                    p_leg.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p_leg.add_run(legenda)
                    run.bold = True
                    run.font.size = Pt(11)
                    try:
                        p_img = doc.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_img = p_img.add_run()
                        run_img.add_picture(caminho_img, width=Inches(5.8))
                    except Exception as ex:
                        print(f"Erro ao adicionar imagem ao docx: {ex}")
                        p_img.clear()
                        p_img.add_run(f"[Imagem indisponivel]")
                    doc.add_paragraph()
                    inseriu = True

            if not inseriu:
                doc.add_paragraph(
                    "Nenhum grafico disponivel."
                )

            # ── 5. METODOLOGIA ─────────────────────────────────────────────────
            doc.add_page_break()
            doc.add_heading("4. Metodologia", level=1)

            metodologias = [
                ("Indice PediRato",
                 "Razao entre contribuicoes ofensivas (Gols + Assistencias) e o valor de mercado "
                 "do atleta em milhoes de euros. Mede quantas contribuicoes ofensivas o clube obtem "
                 "por cada milhao investido. Valores acima de 1.5 indicam alta eficiencia (Pedigree); "
                 "abaixo de 0.5 indicam baixa eficiencia (PediRato)."),

                ("Hype Index (Frente 4)",
                 "Regressao linear simples com X = Performance (Gols + Assists) e Y = Valor de Mercado. "
                 "O residuo de cada jogador mede o quanto seu preco se afasta do esperado pela reta. "
                 "Residuo > +1 desvio padrao: PediRato (caro para o que entrega). "
                 "Residuo < -1 desvio padrao: Pedigree (barato para o que entrega)."),

                ("Espelho Pedigree — Similaridade Cosseno (Frente 5)",
                 "Normaliza as features dos jogadores (Gols, Assists, Valor) com MinMaxScaler e "
                 "calcula a similaridade cosseno entre o jogador de referencia e o pool disponivel. "
                 "Filtra candidatos com valor abaixo de 60% do referencial e retorna os 3 mais "
                 "similares estatisticamente, visualizados em Grafico de Radar."),
            ]

            for titulo_met, descricao in metodologias:
                doc.add_heading(titulo_met, level=2)
                doc.add_paragraph(descricao)

            # ── SALVAR ─────────────────────────────────────────────────────────
            dir_nome = os.path.dirname(filename)
            if dir_nome:
                os.makedirs(dir_nome, exist_ok=True)

            doc.save(filename)

            # Limpar arquivos de imagens temporários
            for caminho_img, _ in graficos_dinamicos:
                if os.path.exists(caminho_img):
                    try:
                        os.remove(caminho_img)
                    except Exception:
                        pass

            print(f"Relatorio Word salvo: {filename}")
            return filename

        except ImportError:
            print("python-docx nao instalado. Execute: pip install python-docx")
            return None
        except Exception as e:
            print(f"Erro ao gerar relatorio Word: {e}")
            return None
